import tkinter as tk
from tkinter import ttk
from tkinter.messagebox import showinfo
from docx import Document
import os

# root window
root = tk.Tk()
root.geometry('400x450')
root.resizable(False, False)
root.title('صدور قرارداد')

# store email address and password
name = tk.StringVar()
father = tk.StringVar()
date = tk.StringVar()
shenas = tk.StringVar()
meli = tk.StringVar()
lisans = tk.StringVar()
maharat = tk.StringVar()
adress = tk.StringVar()


def main():

    template_file_path = '1401.docx'
    output_file_path = name.get()+'.docx'

    variables = {
        "name": name.get(),
        "father": father.get(),
        "date": datee.get(),
        "lisans": lisans.get(),
        "meli": meli.get(),
        "shenas": shenas.get(),
        "maharat": maharat.get(),
        "adress": adress.get(),


    }

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)
    msg = f'قرارداد: {name.get() }صادر شد'
    showinfo(
        title='Information',
        message=msg
    )
    name_entry.delete(0, 1000)
    father_entry.delete(0, 1000)
    date_entry.delete(0, 1000)
    shenas_entry.delete(0, 1000)
    meli_entry.delete(0, 1000)
    lisans_entry.delete(0, 1000)
    maharat_entry.delete(0, 1000)
    adress_entry.delete(0, 1000)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)
                

def del_input():
    name_entry.delete(0, 1000)
    father_entry.delete(0, 1000)
    date_entry.delete(0, 1000)
    shenas_entry.delete(0, 1000)
    meli_entry.delete(0, 1000)
    lisans_entry.delete(0, 1000)
    maharat_entry.delete(0, 1000)
    adress_entry.delete(0, 1000)
    

# Sign in frame
signin = ttk.Frame(root)
signin.pack(padx=10, pady=10, fill='x', expand=True)


# email
name_label = ttk.Label(signin, text="نام و نام خانوادگی:")
name_label.pack(fill='x', expand=True)

name_entry = ttk.Entry(signin, textvariable=name)
name_entry.pack(fill='x', expand=True)
name_entry.focus()

# father
father_label = ttk.Label(signin, text="نام پدر:")
father_label.pack(fill='x', expand=True)

father_entry = ttk.Entry(signin, textvariable=father)
father_entry.pack(fill='x', expand=True)

# datee
date_label = ttk.Label(signin, text="تاریخ تولد:")
date_label.pack(fill='x', expand=True)

date_entry = ttk.Entry(signin, textvariable=date)
date_entry.pack(fill='x', expand=True)

# shenas
shenas_label = ttk.Label(signin, text="شماره شناسنامه:")
shenas_label.pack(fill='x', expand=True)

shenas_entry = ttk.Entry(signin, textvariable=shenas)
shenas_entry.pack(fill='x', expand=True)

# meli
meli_label = ttk.Label(signin, text="شماره ملی:")
meli_label.pack(fill='x', expand=True)

meli_entry = ttk.Entry(signin, textvariable=meli)
meli_entry.pack(fill='x', expand=True)

# lisans
lisans_label = ttk.Label(signin, text="مدرک تحصیلی:")
lisans_label.pack(fill='x', expand=True)

lisans_entry = ttk.Entry(signin, textvariable=lisans)
lisans_entry.pack(fill='x', expand=True)

# maharat
maharat_label = ttk.Label(signin, text="مهارت:")
maharat_label.pack(fill='x', expand=True)

maharat_entry = ttk.Entry(signin, textvariable=maharat)
maharat_entry.pack(fill='x', expand=True)

# adress
adress_label = ttk.Label(signin, text="آدرس:")
adress_label.pack(fill='x', expand=True)

adress_entry = ttk.Entry(signin, textvariable=adress)
adress_entry.pack(fill='x', expand=True)



# login button
login_button = ttk.Button(signin, text="ثبت", command=main)
login_button.pack(fill='x', expand=True, pady=10)

# login button
del_button = ttk.Button(signin, text="پاک کردن", command=del_input)
del_button.pack(fill='x', expand=True, pady=10)


root.mainloop()